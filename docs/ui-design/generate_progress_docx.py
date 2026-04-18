"""Generate the AutoR Studio UI progress document as a .docx.

Run with:

    python docs/ui-design/generate_progress_docx.py

The script writes ``docs/ui-design/AutoR-Studio-UI-Progress.docx`` next to
itself and expects the screenshots at
``docs/ui-design/assets/screenshots/``.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


SCREENSHOT_DIR = Path(__file__).resolve().parent / "assets" / "screenshots"
OUTPUT = Path(__file__).resolve().parent / "AutoR-Studio-UI-Progress.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x18, 0x22, 0x2D)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x18, 0x22, 0x2D)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def add_screenshot(doc: Document, filename: str, caption: str) -> None:
    path = SCREENSHOT_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[missing screenshot: {filename}]")
        return
    doc.add_picture(str(path), width=Inches(6.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure. {caption}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x5B, 0x68, 0x75)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("AutoR Studio — UI Progress Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_paragraph(
        doc,
        "Snapshot of the AutoR Studio frontend on branch `kunlun/ui-initial-design` "
        "as of 2026-04-18. This document summarizes what is shipping, the "
        "NotebookLM-style view that was just added, and the concrete TODOs left "
        "for the next contributor.",
    )

    # ----------------------------------------------------------------- Summary
    add_heading(doc, "1. Where the UI stands", level=1)
    add_paragraph(
        doc,
        "The Studio now ships a working local shell served by `python studio.py "
        "--repo-root . --port 8765`. It runs against real runs on disk — no "
        "mocked data. The top-level structure is:",
    )
    add_bullet(doc, "Hub — projects grid with live run status.")
    add_bullet(doc, "Workspace — per-run view with a page nav.")
    add_bullet(
        doc,
        "Workspace pages: Overview, Human Review, Versions, Notebook. "
        "(The old Files and Paper pages were folded into Notebook.)",
    )
    add_paragraph(
        doc,
        "Every page reads live state from `runs/<run_id>/*` so the two views "
        "(existing 6-stage pipeline UI and the new Notebook view) stay in sync "
        "without a new event bus — the 3-second poll on the workspace view "
        "already catches cross-view changes.",
    )

    # ---------------------------------------------------------------- Hub
    add_heading(doc, "2. Project Hub", level=1)
    add_paragraph(
        doc,
        "Starting surface: list of research projects with latest-run status and "
        "a `New Project` form that seeds both the project record and an initial "
        "run in one click.",
    )
    add_screenshot(doc, "01-hub-view.png", "Project Hub with live run chips and the New Project form.")

    # ---------------------------------------------------------------- Overview
    add_heading(doc, "3. Workspace — Overview", level=1)
    add_paragraph(
        doc,
        "Overview is the pipeline snapshot: stage strip, run progress, activity "
        "feed, and artifact counts. This is the human-friendly supervisor view "
        "of the current run.",
    )
    add_screenshot(doc, "02-workspace-overview.png", "Workspace Overview with live stage strip and run progress.")

    # ---------------------------------------------------------------- Review
    add_heading(doc, "4. Workspace — Human Review", level=1)
    add_paragraph(
        doc,
        "The review page drives the human-in-the-loop gate: a hero block calls "
        "out the stage being reviewed, an inline feedback field + `Approve` / "
        "`Send Feedback & Re-run` buttons are the two primary actions, and a "
        "live Progress Monitor panel shows runner events as they stream in.",
    )
    add_screenshot(doc, "03-workspace-review.png", "Human Review with hero, feedback, and live runner activity.")

    # ---------------------------------------------------------------- Notebook
    add_heading(doc, "5. Workspace — Notebook (new)", level=1)
    add_paragraph(
        doc,
        "The Notebook is a NotebookLM-inspired 3-column surface built around "
        "a real Claude Code instance running inside the run's working directory. "
        "It exists alongside the 6-stage pipeline UI — the two views share the "
        "same on-disk state.",
    )
    add_paragraph(doc, "Columns, left to right:")
    add_bullet(doc, "Sources (left): project thesis, run metadata, stage status, workspace file tree, paper artifacts.")
    add_bullet(doc, "Claude Code (center): conversation with the coordinator, streaming tool calls and thinking events inline.")
    add_bullet(doc, "Viewer (right): clicking anything on the left (or any file-targeted tool chip in the middle) opens that file here. Markdown is rendered, PDFs are embedded, everything else falls back to monospace.")
    add_screenshot(doc, "05-workspace-notebook.png", "Notebook view with sources, CC conversation, and viewer.")

    add_heading(doc, "5.1 How the center panel works", level=2)
    add_paragraph(
        doc,
        "Each user message spawns a one-shot `claude -p --output-format "
        "stream-json --resume <session_id>` subprocess with the run directory "
        "as cwd. The backend streams each stream-json event over SSE to the "
        "browser, appends everything to `runs/<run_id>/notebook/transcript.jsonl`, "
        "and on the very first message captures the session id into "
        "`runs/<run_id>/notebook/session.json` so subsequent messages resume "
        "the same CC conversation.",
    )
    add_paragraph(
        doc,
        "Events render as: user bubbles, assistant markdown blocks, "
        "collapsible `Thinking…` disclosures, tool-use chips with clickable "
        "file targets, and a final `result` footer. System hook spam is "
        "dropped from the visible chat (still kept in the transcript).",
    )

    add_heading(doc, "5.2 Versions", level=2)
    add_paragraph(
        doc,
        "Read-only derived checkpoints based on the run manifest + approved "
        "stages. Named milestones, restore, and branch actions are not yet "
        "implemented.",
    )
    add_screenshot(doc, "04-workspace-versions.png", "Versions page — derived checkpoints, details, and trace timeline.")

    # ---------------------------------------------------------------- Backend
    add_heading(doc, "6. Backend endpoints used by the UI", level=1)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["GET /api/projects/overview", "Project grid with live-run status"],
            ["POST /api/projects", "Create project"],
            ["POST /api/projects/{id}/runs/start", "Start a new run"],
            ["GET /api/runs/{id}", "Run summary + stages"],
            ["GET /api/runs/{id}/stages/{slug}", "Stage markdown"],
            ["POST /api/runs/{id}/stages/{slug}/approve", "Approve stage"],
            ["POST /api/runs/{id}/stages/{slug}/feedback", "Re-run with feedback"],
            ["GET /api/runs/{id}/files/tree", "Workspace file tree"],
            ["GET /api/runs/{id}/files/content", "Read-only file preview"],
            ["GET /api/runs/{id}/paper", "Manuscript summary + sections"],
            ["GET /api/runs/{id}/paper/pdf", "Compiled PDF bytes"],
            ["GET /api/runs/{id}/history", "Derived checkpoints + trace"],
            ["POST /api/notebook/stream", "SSE stream of Claude Code events (new)"],
            ["GET /api/notebook/transcript", "Persisted notebook transcript (new)"],
            ["POST /api/notebook/reset", "Clear notebook session for a run (new)"],
        ],
    )

    # ---------------------------------------------------------------- TODOs
    add_heading(doc, "7. Open TODOs for the next contributor", level=1)
    add_paragraph(doc, "Pick any of these off; each one is sized for a single PR.")

    add_heading(doc, "7.1 Notebook view polish", level=2)
    add_bullet(doc, "Resizable columns: drag handles between the three panels, width persisted to localStorage.")
    add_bullet(doc, "Inline file links in assistant markdown (`[method.tex](...)`) that open in the right-pane viewer via `data-notebook-rel` (scaffold already in notebook.js).")
    add_bullet(doc, "Input affordances: slash-command palette for common AutoR actions (`/approve`, `/rerun-stage`, `/compile`).")
    add_bullet(doc, "Ctrl/Cmd+K shortcut to focus the input from anywhere in the workspace.")
    add_bullet(doc, "Toggle to show/hide hook events in the chat (currently hidden by default — keeps transcript readable).")

    add_heading(doc, "7.2 Sync & robustness", level=2)
    add_bullet(doc, "SSE keepalive ping every 15s so Chrome/proxies don't kill idle connections mid-tool-use.")
    add_bullet(doc, "When CC edits files, have the Overview page auto-bump its stage strip so the user can see the effect without a manual refresh (currently arrives at the next 3s poll).")
    add_bullet(doc, "Race guard: if the user approves a stage from the 6-stage view while CC is holding a reference to that stage, explicitly note it in the next assistant turn.")
    add_bullet(doc, "Browser hard-refresh during a running stream: on reconnect, resume from the transcript and mark the interrupted turn visibly.")

    add_heading(doc, "7.3 Safety & permissions", level=2)
    add_bullet(doc, "Current default is `--permission-mode bypassPermissions`. Offer an `acceptEdits`-only mode behind a toggle for users who want CC to stop at destructive file/bash actions.")
    add_bullet(doc, "Add a configurable allowlist / denylist for Bash commands CC can issue from the Notebook (e.g. deny `rm -rf`, allow `python main.py`, `latex`).")
    add_bullet(doc, "Audit log: keep a per-run summary of the commands CC executed on behalf of the user, surfaced in the Versions timeline.")

    add_heading(doc, "7.4 Write actions missing from the 6-stage UI", level=2)
    add_bullet(doc, "Named milestone persistence (currently `Versions` is read-only).")
    add_bullet(doc, "Restore / branch write actions from the browser.")
    add_bullet(doc, "In-browser file editing (the Notebook viewer is still read-only).")
    add_bullet(doc, "Compile actions on the manuscript artifacts (trigger a LaTeX build from the UI).")
    add_bullet(doc, "VS Code / Overleaf handoff buttons (the Notebook viewer currently shows them disabled as planned pills).")

    add_heading(doc, "7.5 Tests", level=2)
    add_bullet(doc, "Unit tests for `src/backend/notebook.py` session+transcript round-trip against a tmp run dir.")
    add_bullet(doc, "Integration test: stub `claude` binary that echoes predetermined stream-json events, assert SSE event ordering and `done` framing.")
    add_bullet(doc, "UI smoke via Playwright: open hub → open workspace → switch to Notebook → verify the three columns render for a seeded run.")

    # ---------------------------------------------------------------- Run locally
    add_heading(doc, "8. Running locally", level=1)
    add_paragraph(
        doc,
        "From the repo root:",
    )
    code = doc.add_paragraph()
    code_run = code.add_run("python studio.py --repo-root . --port 8765\n# then open http://127.0.0.1:8765 in a browser")
    code_run.font.name = "JetBrains Mono"
    code_run.font.size = Pt(10)
    code.paragraph_format.left_indent = Inches(0.3)

    add_paragraph(
        doc,
        "The Notebook view needs the `claude` CLI on PATH (verified with "
        "`claude --version`). Session state for a run lives at "
        "`runs/<run_id>/notebook/`; delete that directory to start fresh.",
    )

    # ---------------------------------------------------------------- Links
    add_heading(doc, "9. Further reading", level=1)
    add_bullet(doc, "docs/ui-design/README.md — original design pass.")
    add_bullet(doc, "docs/superpowers/specs/2026-04-18-notebook-view-design.md — this session's Notebook design spec.")
    add_bullet(doc, "src/backend/notebook.py — subprocess wrapper + transcript I/O.")
    add_bullet(doc, "src/frontend/static/notebook.js — Notebook client.")

    doc.save(str(OUTPUT))
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    main()
